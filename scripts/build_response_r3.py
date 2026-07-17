"""Build a formatted DOCX point-by-point response from the repository markdown text."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


BLUE = RGBColor(31, 78, 121)
RED = RGBColor(192, 0, 0)

REVISED_ABSTRACT = (
    "Fracture rehabilitation must adapt as pain, function and warning signs change across healing phases, but many digital tools provide static "
    "information and cannot coordinate exercise selection, symptom assessment, progress monitoring and safety escalation. To address this gap, we "
    "developed FractureAgent, a reasoning-and-action (ReAct)-style system that couples a domain-adapted Qwen3.5-9B backbone with five typed "
    "rehabilitation tools and a deterministic safety gate to provide phase-aware support. The model was adapted using quantized low-rank adaptation "
    "(QLoRA) supervised instruction tuning on 18,742 synthetic rehabilitation dialogues and tool-use traces derived from open-access clinical and "
    "patient-education sources. We evaluated FractureAgent using automated functional metrics, expert clinical ratings and 210 simulated-patient "
    "scenarios spanning six fracture types and three rehabilitation phases. It achieved a task completion rate of 91.4%, a mean clinical-appropriateness "
    "score of 4.21/5.00, pain-assessment concordance of 0.873, exercise-appropriateness of 0.896 and complication-detection sensitivity of 0.843, "
    "outperforming the evaluated baselines, including the unfine-tuned Qwen3.5-9B backbone. These findings indicate that FractureAgent can coordinate "
    "multiple rehabilitation tasks and provide adaptive support in simulated settings, addressing the technical limitations of static digital guidance; "
    "however, they do not establish clinical effectiveness or readiness for deployment."
)

REVISED_CODE_AVAILABILITY = (
    "Code availability. The public repository at https://github.com/hangcao87/FractureAgent contains the data-processing scripts, versioned training "
    "prompts, tool schemas and implementations, deterministic safety gate, ms-swift QLoRA configuration and launchers, evaluation code, synthetic "
    "fixtures and non-sensitive run-record templates. The curated training corpus is not publicly distributed because the source materials remain subject "
    "to their original providers' terms and the assembled corpus is being retained for ongoing research, intellectual-property protection and planned "
    "grant-supported work. The trained FractureAgent weights and LoRA checkpoints are not publicly released because follow-on validation and "
    "intellectual-property assessment are ongoing and the resulting model forms part of the same planned research programme. No patient-level data or "
    "private credentials are included. Users may reproduce the disclosed workflow with a legally obtained local base model and their own authorized data; "
    "reasonable requests concerning derived data may be directed to hangcao87@163.com, subject to source licences, author approval and applicable "
    "institutional requirements."
)


def set_font(run, name="Arial", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, before=0, after=6, line=1.08):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_font(r)
    return p


def add_labelled(doc, label, text, color=BLUE, italic=False):
    p = add_para(doc, after=6)
    r = p.add_run(label)
    set_font(r, bold=True, italic=italic, color=color)
    r = p.add_run(text)
    set_font(r, italic=italic, color=color)
    return p


def add_revised_text(doc, label, text):
    p = add_para(doc, before=2, after=9, line=1.05)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    r = p.add_run(label)
    set_font(r, bold=True, color=RED)
    r = p.add_run(text)
    set_font(r, italic=True)
    return p


def build(out: str):
    root = Path(__file__).resolve().parents[1]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(10.5)

    p = add_para(doc, before=0, after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Point-by-Point Response to Reviewers")
    set_font(r, size=16, bold=True, color=BLUE)
    p = add_para(doc, before=0, after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FractureAgent: A Multi-Tool LLM-Based Intelligent Agent for Personalized Fracture Rehabilitation Management")
    set_font(r, size=11.5, bold=True)
    p = add_para(doc, before=0, after=12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Journal: Scientific Reports    Decision: Minor Revision    Manuscript ID: 0d10e894-9f9f-41a6-9d8e-f7f456b33960")
    set_font(r, size=9, color=BLUE)
    add_para(doc, "Dear Editor and Reviewers,", after=8)
    intro = "We sincerely thank the Editor and the Reviewers for their careful assessment and constructive suggestions. We have revised the manuscript and the accompanying research repository to improve abstract readability, code transparency and reproducibility. The revised manuscript is supplied as a tracked-change file. Our point-by-point responses are provided below."
    add_para(doc, intro, after=10)

    add_labelled(doc, "Response to Reviewer 1", "", color=BLUE)
    add_labelled(doc, "Reviewer comment: ", "Thank you.", color=BLUE, italic=True)
    add_labelled(doc, "Response: ", "We thank the Reviewer for the positive assessment.")

    add_labelled(doc, "Response to Reviewer 2", "", color=BLUE)
    add_labelled(doc, "Reviewer comment: ", "The authors have matched all the previously highlighted points. Hence, it is now worthy of publication and I hope it will reach the widest audience possible. https://github.com/hangcao87/FractureAgent", color=BLUE, italic=True)
    add_labelled(doc, "Response: ", "We sincerely thank the Reviewer for the positive assessment and for recognizing the revisions made in the previous round. The public main branch contains the data-processing code, versioned prompts, tool schemas and implementations, deterministic safety gate, ms-swift QLoRA training configuration and launchers, evaluation code, synthetic fixtures and non-sensitive reproducibility-record templates. The curated training corpus and trained model weights are not included for the source-licence, ongoing-validation, planned-research and intellectual-property reasons detailed in our response to Comment 4.2. The repository is available at https://github.com/hangcao87/FractureAgent.")

    doc.add_page_break()
    add_labelled(doc, "Response to Reviewer 4", "", color=BLUE)
    add_labelled(doc, "Comment 4.1 - Reviewer comment: ", "The abstract would benefit from revision to better align with the journal's format and improve readability. The current abstract is structured using subheadings (Background, Methods, Results, Conclusions) and is relatively lengthy. Consider converting it into a single unstructured paragraph and reducing its length by focusing on the most essential methodological details and key findings.", color=BLUE, italic=True)
    add_labelled(doc, "Response: ", "We agree with the Reviewer. The original abstract did not state the clinical-support gap and the design rationale sufficiently clearly. We replaced the four-part structured abstract with a single unstructured paragraph and reduced it from 440 to 184 words, within the journal's 200-word limit. The revision now follows a problem-gap-method-evidence-boundary sequence: it first explains why static digital guidance cannot coordinate changing rehabilitation needs, then shows how the five-tool ReAct architecture and deterministic safety gate address that gap, reports the principal quantitative findings and limits the conclusion to simulated technical feasibility. The complete revised abstract is reproduced below.")
    add_revised_text(doc, "Revised abstract (complete text): ", REVISED_ABSTRACT)
    add_labelled(doc, "Manuscript location: ", "Abstract.", color=RED)
    add_labelled(doc, "Comment 4.2 - Reviewer comment: ", "To enhance transparency and reproducibility, please consider adding a dedicated Code Availability statement. The manuscript provides substantial methodological detail; however, it is currently unclear whether the source code, training scripts, tool implementations, prompts, or model checkpoints are publicly available. If these resources cannot be shared, the authors should explicitly state this and provide the reason.", color=BLUE, italic=True)
    add_labelled(doc, "Response: ", "We agree that the original manuscript did not make the release boundary sufficiently clear. We added a dedicated Code Availability paragraph under Declarations. The public repository contains the data-processing code, versioned prompts, tool schemas and implementations, deterministic safety gate, ms-swift QLoRA configuration and launchers, evaluation code, synthetic fixtures and non-sensitive run-record templates. The curated training corpus is not redistributed because its source materials remain subject to the original providers' terms and the assembled corpus is being retained for ongoing research, intellectual-property protection and planned grant-supported work. The trained model weights and LoRA checkpoints are not released because follow-on validation and intellectual-property assessment remain ongoing and the resulting model is part of the same planned research programme. No patient-level data or private credentials are distributed. To remove a direct contradiction, we also revised only the sentence in the Conclusion that had stated that the model weights would be released openly. The complete Code Availability statement is reproduced below.")
    add_revised_text(doc, "Revised Code Availability statement (complete text): ", REVISED_CODE_AVAILABILITY)
    add_labelled(doc, "Manuscript location: ", "Conclusion, Section 8; Declarations, Code availability.", color=RED)

    add_labelled(doc, "Manuscript change checklist", "", color=BLUE)
    for item in [
        "Abstract: converted the structured Background/Methods/Results/Conclusions format into a 184-word unstructured paragraph and reproduced the complete revision above.",
        "Declarations: added a dedicated Code Availability statement covering data processing, prompts, tools, ms-swift training, evaluation and run records, and stated why the curated corpus and trained weights/checkpoints are not public.",
        "Conclusion: revised only the sentence that had stated that model weights would be released openly, so that it matches the Code Availability statement.",
        "Repository: provides ms-swift installation, data export, QLoRA SFT, inference, evaluation and run-record workflows; no training data or checkpoints are included.",
    ]:
        p = add_para(doc, after=3)
        r = p.add_run("- " + item)
        set_font(r)
    add_para(doc, "Sincerely,", before=12, after=3)
    add_para(doc, "Dr Xiao Ouyang and Dr Weijuan Gong", after=2)
    add_para(doc, "Corresponding authors, on behalf of all authors", after=0)
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", required=True)
    build(ap.parse_args().out)
